from fastapi import APIRouter, HTTPException
from pydantic import BaseModel
from typing import List, Dict, Any, Optional
import pandas as pd
import io
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

router = APIRouter()

class ReportRequest(BaseModel):
    instrument_name: str
    reliability: Optional[Dict[str, Any]] = None
    efa: Optional[Dict[str, Any]] = None
    ai_rel_report: Optional[str] = None
    ai_efa_report: Optional[str] = None

@router.post("/docx")
async def export_docx(req: ReportRequest):
    try:
        doc = Document()
        
        # Style
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Heading
        title_text = f"Reporte de Análisis Psicométrico: {req.instrument_name or 'Instrumento Sin Nombre'}"
        title = doc.add_heading(title_text, 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph("Este reporte ha sido generado automáticamente para fines académicos.")

        # 1. Confiabilidad
        doc.add_heading("1. Análisis de Confiabilidad", level=1)
        
        if req.ai_rel_report:
            doc.add_heading("Interpretación del Experto (IA)", level=2)
            doc.add_paragraph(req.ai_rel_report)

        if req.reliability:
            doc.add_heading("Matriz de Resultados Numéricos", level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Medida de Consistencia'
            hdr_cells[1].text = 'Valor Obtenido'
            hdr_cells[2].text = 'Interpretación'
            
            # Cronbach
            res = req.reliability.get("cronbach", {})
            if res:
                row = table.add_row().cells
                row[0].text = "Alfa de Cronbach"
                row[1].text = str(res.get("alpha", "N/A"))
                row[2].text = "Considerada alta" if (res.get("alpha", 0) or 0) >= 0.7 else "Requiere revisión"
            
            # Omega
            res_o = req.reliability.get("mcdonald_omega", {})
            if res_o and not res_o.get("error"):
                row_o = table.add_row().cells
                row_o[0].text = "Omega de McDonald"
                row_o[1].text = str(res_o.get("omega", "N/A"))
                row_o[2].text = "Consistente" if (res_o.get("omega", 0) or 0) >= 0.7 else "Baja"

            doc.add_heading("Detalles del Procedimiento Matemático", level=2)
            
            if res and "item_stats" in res:
                doc.add_heading("Procedimiento: Alfa de Cronbach", level=3)
                n_items = res.get('n_items', 0)
                
                doc.add_paragraph(f"1) Número de ítems (k) = {n_items}")
                
                # Medias y varianzas
                doc.add_paragraph("2) Estadísticos descriptivos por ítem:")
                table_stats = doc.add_table(rows=1, cols=3)
                table_stats.style = 'Table Grid'
                hdr_stats = table_stats.rows[0].cells
                hdr_stats[0].text = 'Ítem'
                hdr_stats[1].text = 'Media'
                hdr_stats[2].text = 'Desviación Estándar'
                
                for item in res.get("item_stats", []):
                    row = table_stats.add_row().cells
                    row[0].text = str(item.get("item", ""))
                    row[1].text = f"{item.get('mean', 0):.3f}"
                    row[2].text = f"{item.get('std_dev', 0):.3f}"
                
                doc.add_paragraph("\n3) Parámetros de la fórmula:")
                doc.add_paragraph(f"• ∑Sᵢ² = {res.get('sum_item_variances', 'N/A')} (Suma de las varianzas individuales)")
                doc.add_paragraph(f"• Sₜ² = {res.get('total_variance', 'N/A')} (Varianza del puntaje total)")
                
                doc.add_paragraph("\n4) Sustitución de la fórmula:")
                k = res.get('n_items', 'N/A')
                sum_var = res.get('sum_item_variances', 'N/A')
                tot_var = res.get('total_variance', 'N/A')
                alpha = res.get('alpha', 'N/A')
                doc.add_paragraph(f"α = [{k} / ({k} - 1)] × [1 - ({sum_var} / {tot_var})] = {alpha}")
                
            if res_o and not res_o.get("error") and "item_stats" in res_o:
                doc.add_heading("Procedimiento: Omega de McDonald", level=3)
                n_items = res.get('n_items', 'N/A') if res else len(res_o.get("item_stats", []))
                n_cases = res_o.get('n_cases', 'N/A')
                doc.add_paragraph(f"1) Número de ítems (k) = {n_items}")
                doc.add_paragraph(f"Número de participantes (n) = {n_cases}\n")
                
                doc.add_paragraph("2) Matriz de Correlaciones (Resumen de consistencia entre ítems)")
                correlations = res_o.get("correlations", {})
                if correlations:
                    cols = list(correlations.keys())
                    table_corr = doc.add_table(rows=1, cols=len(cols) + 1)
                    table_corr.style = 'Table Grid'
                    hdr_corr = table_corr.rows[0].cells
                    hdr_corr[0].text = "Ítem"
                    for idx, c in enumerate(cols):
                        hdr_corr[idx + 1].text = str(c)
                    
                    for r_key, row_dict in correlations.items():
                        row = table_corr.add_row().cells
                        row[0].text = str(r_key)
                        for idx, c in enumerate(cols):
                            val = row_dict.get(c, 0)
                            row[idx + 1].text = str(val) if val != 0 else "0.0"
                
                doc.add_paragraph("\n3) Extraer factor común y varianza de error")
                p_note = doc.add_paragraph()
                p_note.add_run("Nota: ").bold = True
                p_note.add_run("Análisis factorial de un solo factor (Principal Axis Factoring).")
                
                table_fa = doc.add_table(rows=1, cols=3)
                table_fa.style = 'Table Grid'
                hdr_fa = table_fa.rows[0].cells
                hdr_fa[0].text = 'Ítem'
                hdr_fa[1].text = 'Carga Factorial (λ)'
                hdr_fa[2].text = 'Varianza de Error (θ)'
                
                for item in res_o.get("item_stats", []):
                    row = table_fa.add_row().cells
                    row[0].text = str(item.get("item", ""))
                    row[1].text = f"{item.get('loading', 0):.3f}"
                    row[2].text = f"{item.get('uniqueness', 0):.3f}"

                doc.add_paragraph("\n4) Sumatorias:")
                doc.add_paragraph(f"• Suma de cargas (∑λ) = {res_o.get('sum_loadings', 'N/A')}")
                doc.add_paragraph(f"• Suma de errores (∑θ) = {res_o.get('sum_uniqueness', 'N/A')}")
                doc.add_paragraph(f"• (∑λ)² = {res_o.get('sum_loadings_sq', 'N/A')}")
                
                doc.add_paragraph("\n5) Fórmula y Sustitución:")
                sum_sq = res_o.get('sum_loadings_sq', 'N/A')
                sum_u = res_o.get('sum_uniqueness', 'N/A')
                omega = res_o.get('omega', 'N/A')
                doc.add_paragraph(f"ω = {sum_sq} / ({sum_sq} + {sum_u}) = {omega}")
                doc.add_paragraph(f"\nResultado final: ω = {omega}")

        # 2. Análisis Factorial
        doc.add_heading("2. Análisis Factorial Exploratorio (EFA)", level=1)
        
        if req.ai_efa_report:
            doc.add_heading("Interpretación de la Estructura (IA)", level=2)
            doc.add_paragraph(req.ai_efa_report)
            
        if req.efa:
            doc.add_heading("Adecuación Muestral", level=2)
            ade = req.efa.get("adequacy", {})
            kmo = ade.get("kmo", {})
            bart = ade.get("bartlett", {})
            
            p = doc.add_paragraph()
            p.add_run(f"- KMO (Kaiser-Meyer-Olkin): ").bold = True
            p.add_run(f"{kmo.get('value')} ({kmo.get('interpretation')})\n")
            p.add_run(f"- Prueba de Bartlett: ").bold = True
            p.add_run(f"Chi2={bart.get('chi_square')}, p-valor={bart.get('p_value')}\n")
            
            # Loadings Matrix
            loadings = req.efa.get("loadings", [])
            if loadings:
                doc.add_heading("Matriz de Cargas Factoriales", level=2)
                cols = list(loadings[0].keys())
                table_efa = doc.add_table(rows=1, cols=len(cols))
                table_efa.style = 'Table Grid'
                hdr = table_efa.rows[0].cells
                for i, c in enumerate(cols):
                    hdr[i].text = c
                
                for row_data in loadings:
                    cells = table_efa.add_row().cells
                    for i, c in enumerate(cols):
                        val = row_data[c]
                        if isinstance(val, (int, float)):
                            cells[i].text = f"{val:.3f}"
                        else:
                            cells[i].text = str(val)

        # Save to memory
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        safe_name = (req.instrument_name or "Reporte").replace(" ", "_")
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename=Reporte_{safe_name}.docx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.post("/xlsx")
async def export_xlsx(req: ReportRequest):
    try:
        output = io.BytesIO()
        # Use pandas with xlsxwriter
        with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
            # Summary
            if req.reliability:
                rel_data = {
                    "Métrica": ["Alfa de Cronbach", "Omega de McDonald"],
                    "Valor": [
                        req.reliability.get("cronbach", {}).get("alpha"),
                        req.reliability.get("mcdonald_omega", {}).get("omega")
                    ],
                    "N Ítems": [
                        req.reliability.get("cronbach", {}).get("n_items"),
                        "N/A"
                    ]
                }
                pd.DataFrame(rel_data).to_excel(writer, sheet_name="Confiabilidad", index=False)
            
            # Loadings
            if req.efa and req.efa.get("loadings"):
                pd.DataFrame(req.efa["loadings"]).to_excel(writer, sheet_name="Matriz de Cargas", index=False)
            
            # Eigenvalues (if available)
            if req.efa and req.efa.get("extraction", {}).get("eigenvalues"):
                eig_data = pd.DataFrame({
                    "Componente": range(1, len(req.efa["extraction"]["eigenvalues"]) + 1),
                    "Autovalor (Eigenvalue)": req.efa["extraction"]["eigenvalues"]
                })
                eig_data.to_excel(writer, sheet_name="Autovalores", index=False)
                
        output.seek(0)
        safe_name = (req.instrument_name or "Datos").replace(" ", "_")
        return StreamingResponse(
            output,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename=Matriz_{safe_name}.xlsx"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
